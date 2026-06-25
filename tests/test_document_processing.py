import pymupdf
import pytest
from docx import Document as DocxDocument

from app.rag.chunker import chunk_document
from app.rag.document_loader import (
    DocumentLoadError,
    UnsupportedDocumentTypeError,
    load_document,
    load_uploaded_document,
)


def test_load_txt_document(tmp_path) -> None:
    path = tmp_path / "note.txt"
    path.write_bytes(b"hello\r\nworld\n")

    document = load_document(path)

    assert document.content == "hello\nworld"
    assert document.metadata.source == str(path)
    assert document.metadata.file_name == "note.txt"
    assert document.metadata.file_type == "txt"
    assert document.document_id


def test_load_markdown_document(tmp_path) -> None:
    path = tmp_path / "readme.md"
    path.write_text("# Title\n\nProject notes", encoding="utf-8")

    document = load_document(path)

    assert document.content.startswith("# Title")
    assert document.metadata.file_type == "md"


def test_load_docx_document_extracts_paragraphs_and_tables(tmp_path) -> None:
    path = tmp_path / "resume.docx"
    docx = DocxDocument()
    docx.add_paragraph("个人简历")
    docx.add_paragraph("熟悉 FastAPI、Redis 和 LangGraph。")
    table = docx.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "AI Resume Job Agent"
    docx.save(path)

    document = load_document(path)

    assert document.metadata.source == str(path)
    assert document.metadata.file_name == "resume.docx"
    assert document.metadata.file_type == "docx"
    assert "个人简历" in document.content
    assert "熟悉 FastAPI、Redis 和 LangGraph。" in document.content
    assert "项目 | AI Resume Job Agent" in document.content


def test_load_uploaded_docx_document(tmp_path) -> None:
    path = tmp_path / "resume.docx"
    docx = DocxDocument()
    docx.add_paragraph("求职方向：AI Agent 开发")
    docx.save(path)

    document = load_uploaded_document("resume.docx", path.read_bytes())

    assert document.metadata.source == "upload:resume.docx"
    assert document.metadata.file_type == "docx"
    assert document.content == "求职方向：AI Agent 开发"


def test_load_uploaded_docx_rejects_invalid_bytes() -> None:
    with pytest.raises(DocumentLoadError, match="not a valid .docx"):
        load_uploaded_document("broken.docx", b"not a real docx")


def test_load_docx_rejects_empty_document(tmp_path) -> None:
    path = tmp_path / "empty.docx"
    DocxDocument().save(path)

    with pytest.raises(DocumentLoadError, match="Document is empty"):
        load_document(path)


def test_load_document_rejects_unsupported_suffix(tmp_path) -> None:
    path = tmp_path / "slides.pptx"
    path.write_text("fake pptx content", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentTypeError):
        load_document(path)


def test_load_text_layer_pdf_document(tmp_path) -> None:
    path = tmp_path / "resume.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Resume RAG Agent FastAPI Redis", fontsize=14)
    document.save(path)
    document.close()

    loaded = load_document(path)

    assert loaded.metadata.file_type == "pdf"
    assert loaded.metadata.file_name == "resume.pdf"
    assert "Resume RAG Agent" in loaded.content


def test_load_uploaded_pdf_rejects_invalid_bytes() -> None:
    with pytest.raises(DocumentLoadError, match="not a valid .pdf"):
        load_uploaded_document("broken.pdf", b"this is not a pdf")


def test_chunk_document_splits_text_with_overlap(tmp_path) -> None:
    path = tmp_path / "long.txt"
    path.write_text("abcdefghijklmnopqrstuvwxyz", encoding="utf-8")
    document = load_document(path)

    chunks = chunk_document(document, chunk_size=10, chunk_overlap=2)

    assert [chunk.content for chunk in chunks] == [
        "abcdefghij",
        "ijklmnopqr",
        "qrstuvwxyz",
    ]
    assert [chunk.start_char for chunk in chunks] == [0, 8, 16]
    assert [chunk.end_char for chunk in chunks] == [10, 18, 26]


def test_chunk_document_preserves_metadata_and_chunk_index(tmp_path) -> None:
    path = tmp_path / "knowledge.md"
    path.write_text("agent rag redis", encoding="utf-8")
    document = load_document(path)

    chunks = chunk_document(document, chunk_size=20, chunk_overlap=0)

    assert len(chunks) == 1
    assert chunks[0].chunk_id == f"{document.document_id}:0"
    assert chunks[0].document_id == document.document_id
    assert chunks[0].source == str(path)
    assert chunks[0].chunk_index == 0
    assert chunks[0].metadata["source"] == str(path)
    assert chunks[0].metadata["file_name"] == "knowledge.md"
    assert chunks[0].metadata["file_type"] == "md"
    assert chunks[0].metadata["chunk_index"] == 0


def test_chunk_document_rejects_invalid_overlap(tmp_path) -> None:
    path = tmp_path / "note.txt"
    path.write_text("hello world", encoding="utf-8")
    document = load_document(path)

    with pytest.raises(ValueError):
        chunk_document(document, chunk_size=10, chunk_overlap=10)
