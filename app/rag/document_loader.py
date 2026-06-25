from hashlib import sha256
from io import BytesIO
from pathlib import Path

from app.models.document import Document, DocumentMetadata
from app.rag.pdf_extractor import PdfExtractionError, extract_pdf_text


SUPPORTED_TEXT_SUFFIXES = {".txt", ".md"}
SUPPORTED_DOCX_SUFFIXES = {".docx"}
SUPPORTED_PDF_SUFFIXES = {".pdf"}
SUPPORTED_SUFFIXES = (
    SUPPORTED_TEXT_SUFFIXES | SUPPORTED_DOCX_SUFFIXES | SUPPORTED_PDF_SUFFIXES
)


class UnsupportedDocumentTypeError(ValueError):
    pass


class DocumentLoadError(ValueError):
    pass


def load_document(path: str | Path) -> Document:
    document_path = Path(path)
    suffix = document_path.suffix.lower()

    if suffix not in SUPPORTED_SUFFIXES:
        supported = _format_supported_suffixes()
        raise UnsupportedDocumentTypeError(f"Unsupported document type: {suffix}. Supported: {supported}")

    if suffix in SUPPORTED_DOCX_SUFFIXES:
        content = _load_docx_from_path(document_path)
    elif suffix in SUPPORTED_PDF_SUFFIXES:
        content = _load_pdf_from_path(document_path)
    else:
        content = _normalize_text(document_path.read_text(encoding="utf-8"))
    source = str(document_path)
    return _build_document(
        source=source,
        file_name=document_path.name,
        file_type=suffix.lstrip("."),
        content=content,
    )


def load_uploaded_document(file_name: str, data: bytes) -> Document:
    safe_file_name = Path(file_name).name
    suffix = Path(safe_file_name).suffix.lower()

    if suffix not in SUPPORTED_SUFFIXES:
        supported = _format_supported_suffixes()
        raise UnsupportedDocumentTypeError(f"Unsupported document type: {suffix}. Supported: {supported}")

    if suffix in SUPPORTED_DOCX_SUFFIXES:
        content = _load_docx_from_bytes(data, safe_file_name)
    elif suffix in SUPPORTED_PDF_SUFFIXES:
        content = _load_pdf_from_bytes(data, safe_file_name)
    else:
        try:
            content = _normalize_text(data.decode("utf-8"))
        except UnicodeDecodeError as exc:
            raise DocumentLoadError(f"Document must be UTF-8 text: {safe_file_name}") from exc

    return _build_document(
        source=f"upload:{safe_file_name}",
        file_name=safe_file_name,
        file_type=suffix.lstrip("."),
        content=content,
    )


def _build_document(
    *,
    source: str,
    file_name: str,
    file_type: str,
    content: str,
) -> Document:
    if not content:
        raise DocumentLoadError(f"Document is empty: {source}")

    return Document(
        document_id=_build_document_id(source, content),
        content=content,
        metadata=DocumentMetadata(
            source=source,
            file_name=file_name,
            file_type=file_type,
        ),
    )


def _normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def _build_document_id(source: str, content: str) -> str:
    return sha256(f"{source}\n{content}".encode("utf-8")).hexdigest()[:16]


def _load_docx_from_path(path: Path) -> str:
    try:
        document = _open_docx_document(path)
    except Exception as exc:
        raise DocumentLoadError(f"Document is not a valid .docx file: {path}") from exc
    return _normalize_text(_extract_docx_text(document))


def _load_docx_from_bytes(data: bytes, file_name: str) -> str:
    try:
        document = _open_docx_document(BytesIO(data))
    except Exception as exc:
        raise DocumentLoadError(f"Document is not a valid .docx file: {file_name}") from exc
    return _normalize_text(_extract_docx_text(document))


def _open_docx_document(source):
    from docx import Document as DocxDocument

    return DocxDocument(source)


def _extract_docx_text(document) -> str:
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [_normalize_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _load_pdf_from_path(path: Path) -> str:
    try:
        data = path.read_bytes()
    except OSError as exc:
        raise DocumentLoadError(f"Cannot read PDF file: {path}") from exc
    return _extract_pdf_content(data)


def _load_pdf_from_bytes(data: bytes, file_name: str) -> str:
    return _extract_pdf_content(data)


def _extract_pdf_content(data: bytes) -> str:
    try:
        result = extract_pdf_text(data)
    except PdfExtractionError as exc:
        raise DocumentLoadError(str(exc)) from exc
    return _normalize_text(result.text)


def _format_supported_suffixes() -> str:
    return ", ".join(sorted(SUPPORTED_SUFFIXES))
